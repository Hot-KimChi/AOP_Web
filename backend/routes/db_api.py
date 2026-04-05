from flask import Blueprint, request, jsonify, g, Response
import os, io
from pathlib import Path
from docx import Document
from utils.decorators import handle_exceptions, require_auth, with_db_connection
from utils.error_handler import error_response
from utils.logger import logger
import pandas as pd

db_api_bp = Blueprint("db_api", __name__, url_prefix="/api")


@db_api_bp.route("/insert-sql", methods=["POST"])
@handle_exceptions
@require_auth
@with_db_connection()
def insert_sql_measset():
    data = request.get_json()
    table_name = data.get("table")
    records = data.get("data")
    if not table_name or not records:
        return error_response("Invalid data: table and data fields are required", 400)
    try:
        import json
        from io import StringIO

        json_str = json.dumps(records)
        df = pd.read_json(StringIO(json_str), orient="records")
    except Exception as e:
        logger.error(f"DataFrame conversion error: {str(e)}")
        return error_response("Failed to parse records into DataFrame", 500)
    try:
        g.current_db.insert_data(table_name, df)
        return (
            jsonify({"status": "success", "message": "Data inserted successfully"}),
            200,
        )
    except Exception as e:
        logger.error(f"Data insertion failed: {str(e)}", exc_info=True)
        return error_response(str(e), 500)


@db_api_bp.route("/csv-data", methods=["GET"])
@handle_exceptions
@require_auth
def get_csv_data():
    csv_key = request.args.get("csv_key")
    if not csv_key:
        return error_response("csv_key is required", 400)

    # 경로 탐색 공격 방지: 절대 경로로 변환 후 프로젝트 루트 내에 있는지 검증
    try:
        file_path = Path(csv_key).resolve()
        project_root = Path(os.getcwd()).resolve()
        file_path.relative_to(project_root)
    except ValueError:
        logger.warning(f"Path traversal attempt blocked: {csv_key!r}")
        return error_response("Invalid file path", 400)
    except Exception:
        return error_response("Invalid file path", 400)

    if not file_path.exists():
        return error_response("CSV data not found", 404)
    with open(file_path, "r", encoding="utf-8") as f:
        csv_data = f.read()
    return jsonify({"status": "success", "data": csv_data}), 200


@db_api_bp.route("/get_list_database", methods=["GET"])
@handle_exceptions
@require_auth
def get_list_database():
    raw = os.environ.get("DATABASE_NAME", "")
    databases = [d.strip() for d in raw.split(",") if d.strip()]
    return jsonify({"status": "success", "databases": databases})


@db_api_bp.route("/get_list_table", methods=["GET"])
@handle_exceptions
@require_auth
def get_list_table():
    raw = os.environ.get("SERVER_TABLE_TABLE", "")
    tables = [t.strip() for t in raw.split(",") if t.strip()]
    return jsonify({"status": "success", "tables": tables})


@db_api_bp.route("/get_probes", methods=["GET"])
@handle_exceptions
@require_auth
@with_db_connection()
def get_probes():
    selected_database = request.args.get("database")
    selected_table = request.args.get("table")
    logger.info(f"Database: {selected_database}, Table: {selected_table}")
    allowed_tables = ["Tx_summary", "probe_geo"]
    if selected_table not in allowed_tables:
        return (
            jsonify({"status": "error", "message": "유효하지 않은 테이블 이름입니다"}),
            400,
        )
    query = f"SELECT probeId, probeName FROM {selected_table}"
    df = g.current_db.execute_query(query)
    df["probeId"] = df["probeId"].fillna("empty")
    df_unique = df.drop_duplicates(subset=["probeId", "probeName"])
    df_unique = df_unique.sort_values(by="probeName").reset_index(drop=True)
    df_unique["_id"] = df_unique["probeId"].astype(str) + "_" + df_unique.index.astype(str)
    probes = df_unique[["probeId", "probeName", "_id"]].to_dict("records")
    return jsonify({"status": "success", "probes": probes})


@db_api_bp.route("/get_table_data", methods=["GET"])
@handle_exceptions
@require_auth
@with_db_connection()
def get_table_data():
    selected_database = request.args.get("database")
    selected_table = request.args.get("table")
    logger.info(f"Database: {selected_database}, Table: {selected_table}")
    allowed_tables = ["Tx_summary", "probe_geo", "WCS", "meas_station_setup"]
    if selected_table not in allowed_tables:
        return (
            jsonify({"status": "error", "message": "유효하지 않은 테이블 이름입니다"}),
            400,
        )
    if selected_table == "meas_station_setup":
        df = g.current_db.execute_query(
            f"SELECT measSSId, measComments, probeId, measPersonName, measPurpose, imagingSysSn, probeSn, hydrophId FROM {selected_table} where measPurpose not like '%Beamstyle%' order by measSSId desc"
        )
        data = df.to_dict(orient="records") if df is not None else []
        columns = list(df.columns) if df is not None else []
        return jsonify({"status": "success", "data": data, "columns": columns})
    if selected_table == "Tx_summary":
        query = f"SELECT DISTINCT ProbeID AS probeId, ProbeName AS probeName, Software_version AS software_version FROM {selected_table} ORDER BY software_version DESC"
    elif selected_table == "WCS":
        query = f"SELECT DISTINCT probeId, myVersion FROM {selected_table} ORDER BY myVersion DESC"
    else:
        query = f"SELECT DISTINCT probeId, probeName FROM {selected_table}"
    df = g.current_db.execute_query(query)
    if selected_table == "WCS":
        df = df.dropna(subset=["probeId", "myVersion"])
    else:
        df = df.dropna(subset=["probeId", "probeName"])
    response_data = {
        "status": "success",
        "hasSoftwareData": selected_table == "Tx_summary",
    }
    if selected_table == "WCS":
        df["probeId"] = df["probeId"].astype(str)
        df["myVersion"] = df["myVersion"].astype(str)
        df = df.drop_duplicates(subset=["probeId", "myVersion"]).reset_index(drop=True)
        df["_id"] = "wcs_" + df.index.astype(str)
        response_data["wcsVersions"] = df[["probeId", "myVersion", "_id"]].to_dict("records")
        return jsonify(response_data)
    df_probes = df.drop_duplicates(subset=["probeId", "probeName"])
    df_probes = df_probes.sort_values(by="probeName").reset_index(drop=True)
    df_probes["probeId"] = df_probes["probeId"].astype(str)
    df_probes["probeName"] = df_probes["probeName"].astype(str)
    df_probes["_id"] = df_probes["probeId"] + "_" + df_probes.index.astype(str)
    probes = df_probes[["probeId", "probeName", "_id"]].to_dict("records")
    response_data["probes"] = probes
    if selected_table == "Tx_summary":
        df["software_version"] = df["software_version"].fillna("Empty")
        df["software_version"] = df["software_version"].astype(str)
        df_software = df.drop_duplicates(subset=["software_version"])
        df_software = df_software.sort_values(
            by="software_version", key=lambda x: x.astype(str)
        )
        # probe_software_map: groupby로 O(n²) 중복 탐지 제거
        df_sw_map = df[["probeId", "software_version"]].copy()
        df_sw_map["probeId"] = df_sw_map["probeId"].astype(str)
        df_sw_map = df_sw_map[df_sw_map["software_version"].str.lower() != "empty"]
        probe_software_map = {
            probe_id: [{"softwareVersion": v} for v in group["software_version"].unique()]
            for probe_id, group in df_sw_map.groupby("probeId", sort=False)
        }
        # software 목록 벡터화
        df_sw_filtered = df_software[
            df_software["software_version"].str.lower() != "empty"
        ].reset_index(drop=True)
        df_sw_filtered["_id"] = "sw_version_" + df_sw_filtered.index.astype(str)
        software = df_sw_filtered.rename(
            columns={"software_version": "softwareVersion"}
        )[["softwareVersion", "_id"]].to_dict("records")
        response_data["software"] = software
        response_data["mapping"] = probe_software_map
    return jsonify(response_data)


@db_api_bp.route("/run_tx_compare", methods=["POST"])
@handle_exceptions
@require_auth
@with_db_connection()
def run_tx_compare():
    try:
        if not request.is_json:
            return error_response(
                "요청 형식이 잘못되었습니다. JSON 형식이 필요합니다.", 400
            )
        data = request.get_json()
        required_params = ["probeId", "TxSumSoftware", "wcsSoftware"]
        missing_params = [param for param in required_params if not data.get(param)]
        if missing_params:
            return error_response(
                f"필수 파라미터가 누락되었습니다: {', '.join(missing_params)}", 400
            )
        ssid_temp = data.get("measSSId_Temp")
        ssid_mi = data.get("measSSId_MI")
        ssid_ispta3 = data.get("measSSId_Ispta")
        if not (ssid_temp or ssid_mi or ssid_ispta3):
            return error_response(
                "measSSId_Temp, measSSId_MI, measSSId_Ispta 중 적어도 하나는 입력해야 합니다.",
                400,
            )
        probeid = int(float(data.get("probeId")))
        tx_sw = data.get("TxSumSoftware")
        wcs_sw = data.get("wcsSoftware")
        ssid_temp = None if ssid_temp == "" or ssid_temp is None else ssid_temp
        ssid_mi = None if ssid_mi == "" or ssid_mi is None else ssid_mi
        ssid_ispta3 = None if ssid_ispta3 == "" or ssid_ispta3 is None else ssid_ispta3
        logger.info(
            f"TxCompare 실행 요청: probeId={probeid}, Tx_SW={tx_sw}, WCS_SW={wcs_sw}, SSid_Temp={ssid_temp}, SSid_MI={ssid_mi}, SSid_Ispta3={ssid_ispta3}"
        )
        params = (probeid, tx_sw, wcs_sw, ssid_temp, ssid_mi, ssid_ispta3)
        result_df = g.current_db.execute_procedure("TxCompare", params)
        import numpy as np

        result_df = result_df.replace({np.nan: None})
        if result_df is None or result_df.empty:
            return (
                jsonify(
                    {
                        "status": "success",
                        "message": "비교 보고서 데이터가 없습니다.",
                        "reportData": [],
                    }
                ),
                200,
            )
        report_data = result_df.to_dict(orient="records")
        columns = list(result_df.columns)
        return (
            jsonify(
                {
                    "status": "success",
                    "message": "비교 보고서 데이터를 성공적으로 추출했습니다.",
                    "reportData": report_data,
                    "columns": columns,
                }
            ),
            200,
        )
    except Exception as e:
        logger.error(f"TxCompare 실행 중 오류 발생: {str(e)}", exc_info=True)
        return error_response(
            f"비교 보고서 데이터 추출 중 오류가 발생했습니다: {str(e)}", 500
        )


@db_api_bp.route("/export_table_to_word", methods=["GET"])
@handle_exceptions
@require_auth
@with_db_connection()
def export_table_to_word():
    selected_database = request.args.get("database")
    selected_table = request.args.get("table")
    measSSIds = request.args.get("measSSIds")
    if not selected_database or not selected_table:
        return error_response("database, table 파라미터가 필요합니다", 400)
    # SQL 인젝션 방지: 테이블명 allowlist 검증
    allowed_export_tables = ["SSR_table", "Tx_summary", "probe_geo", "WCS", "meas_station_setup"]
    if selected_table not in allowed_export_tables:
        return (
            jsonify({"status": "error", "message": "유효하지 않은 테이블 이름입니다"}),
            400,
        )
    if measSSIds:
        id_list = [s for s in measSSIds.split(",") if s.isdigit()]
        if not id_list:
            return error_response("measSSIds 파라미터가 올바르지 않습니다", 400)
        id_str = ",".join(id_list)
        query = f"SELECT * FROM {selected_table} WHERE measSSId IN ({id_str})"
    else:
        query = f"SELECT * FROM {selected_table}"
    df = g.current_db.execute_query(query)
    if df is None or df.empty:
        return error_response("해당 테이블에 데이터가 없습니다", 404)
    doc = Document()
    doc.add_heading(f"Table: {selected_table}", 0)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for row_values in df.fillna("").astype(str).values.tolist():
        row_cells = table.add_row().cells
        for i, value in enumerate(row_values):
            row_cells[i].text = value
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"{selected_table}.docx"
    return Response(
        buf.getvalue(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@db_api_bp.route("/get_viewer_data", methods=["GET"])
@handle_exceptions
@require_auth
@with_db_connection()
def get_viewer_data():
    """Database Viewer 팝업 전용: 선택된 테이블의 전체 데이터(TOP 1000)를 반환합니다."""
    selected_database = request.args.get("database")
    selected_table = request.args.get("table")

    if not selected_database or not selected_table:
        return error_response("database, table 파라미터가 필요합니다", 400)

    # SERVER_TABLE_TABLE 환경변수 기준 동적 allowlist (get_list_table과 동기화)
    allowed_tables_raw = os.environ.get("SERVER_TABLE_TABLE", "")
    allowed_tables = [t.strip() for t in allowed_tables_raw.split(",") if t.strip()]
    if selected_table not in allowed_tables:
        return error_response("유효하지 않은 테이블 이름입니다", 400)

    import numpy as np

    df = g.current_db.execute_query(f"SELECT TOP 1000 * FROM {selected_table}")
    if df is None or df.empty:
        return jsonify({"status": "success", "data": [], "columns": []})

    df = df.replace({np.nan: None})
    return jsonify({
        "status": "success",
        "data": df.to_dict(orient="records"),
        "columns": list(df.columns),
    })
